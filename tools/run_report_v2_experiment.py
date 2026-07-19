#!/usr/bin/env python3
"""Local HEAL LLM2/report v2 experiment.

This is intentionally not wired into the production API. It lets us test a
stronger final synthesis prompt against an existing normalized LLM1 CSV and
render the result as a human-readable DOCX.
"""

from __future__ import annotations

import argparse
import csv
import datetime as dt
import importlib.util
import json
import os
import sys
import urllib.error
import urllib.request
from pathlib import Path


OPENAI_RESPONSES_URL = "https://api.openai.com/v1/responses"
REPO_ROOT = Path(__file__).resolve().parents[1]
GLOBAL_SCRIPT = REPO_ROOT / "services" / "heal-global-interpretation" / "interpret_global_profile.py"


def load_global_module():
    spec = importlib.util.spec_from_file_location("heal_global_interpretation", GLOBAL_SCRIPT)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Could not load {GLOBAL_SCRIPT}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def load_env_file(path: Path) -> None:
    if not path.exists():
        return
    for raw_line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip())


def read_csv(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def utc_now() -> str:
    return dt.datetime.now(dt.UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def output_text_from_response(response: dict) -> str:
    texts = []
    for item in response.get("output") or []:
        for content in item.get("content") or []:
            if content.get("type") in {"output_text", "text"} and content.get("text"):
                texts.append(content["text"])
    if texts:
        return "\n".join(texts)
    if response.get("output_text"):
        return str(response["output_text"]).strip()
    raise ValueError("OpenAI response did not contain output text.")


def report_style_guidance() -> str:
    return """

REPORT STRUCTURE AND STYLE ADDENDUM

Apply these rules generically to every case. Do not hard-code specific genes,
rsIDs, modules, or diseases from any example.

1. Reader-first structure:
   - Start with "1. Panorama general del caso".
   - Then "2. Ejes biologicos principales".
   - Then "3. Hallazgos secundarios, limitaciones y pasos de revision".
   - Avoid repeating the same disclaimer in each section. State the non-diagnostic
     boundary once, then use short axis-specific cautions only when needed.

2. Interpretation before evidence:
   - In narrative text, do not interrupt the biological interpretation with long
     parenthetical lists of genes or rsIDs.
   - Do not place gene or rsID lists in the executive summary, main verdict, or
     opening paragraphs. Those sections should name systems, not enumerate genes.
   - Do not use slash-separated gene symbol strings in narrative fields.
   - Gene symbols should appear primarily in the dedicated support fields
     (supporting_genes, supporting_rsids) and in sensitive/conflicting finding
     entries where the exact variant is necessary.
   - In narrative fields, prefer biological labels such as "cytokine signaling",
     "drug-metabolism enzymes", "dopaminergic/serotonergic signaling", or
     "one-carbon metabolism" instead of listing gene symbols.
   - First explain what the axis may mean biologically.
   - At the end of that axis, add a compact support block with:
     "Genes que soportan este eje:" and "rsIDs relevantes:".
   - If evidence is mixed, say that in the support block instead of overloading
     the main paragraph.
   - Axis titles should be short and readable. Avoid long parenthetical gene
     lists in titles.

3. Contextual branching is allowed:
   - You may describe what an axis could reasonably point toward for contextual
     review, using wording such as "podria justificar revisar", "podria orientar
     una conversacion sobre", or "si hay sintomas o antecedentes, podria valer
     la pena discutir".
   - You may name the type of professional or review area when directly supported
     by the axis, such as immunology for immune/inflammatory axes, clinical
     genetics for conflicting variants, pharmacogenomics/pharmacy for medication
     metabolism axes, or nutrition/metabolism for folate/B12/metabolic axes.
   - Do not give treatment, supplement, medication, lifestyle, or diagnostic
     recommendations.
   - For every primary biological axis, provide a useful contextual review
     recommendation. Do not only say "consult a professional". Explain what
     domain could be relevant, when it could matter, who could review it, what it
     should clarify, and what must not be inferred.
   - Allowed examples, generalized:
     * For neurotransmission axes: if there are compatible symptoms or concerns
       about attention, impulsivity, stress regulation, learning, sleep, mood, or
       development, the axis may be worth discussing in a formal
       neurodevelopmental, psychological, psychiatric, neurological, or clinical
       evaluation.
     * For immune/inflammatory axes: if there are allergic, inflammatory,
       autoimmune-like symptoms, recurrent infections, or relevant family
       history, the axis may be worth reviewing with a clinician, immunologist,
       allergist, or relevant specialist.
     * For nutrient/methylation axes: if there are abnormal folate, B12,
       homocysteine, methylation-related labs, medication context, or relevant
       family history, the axis may be worth reviewing with a clinician or
       qualified nutrition/metabolism professional.
     * For PGx axes: if there are current or planned medications, adverse drug
       reactions, or pharmacogenomic questions, the axis may be worth reviewing
       with a pharmacogenomics professional.
     * For sleep/circadian axes: if there are sleep complaints, circadian issues,
       stimulant/caffeine sensitivity, or fatigue concerns, the axis may be worth
       contextual review.
   - Use wording like: "This does not diagnose X. However, if there are symptoms,
     family history, or clinical concerns compatible with X or related
     conditions, this axis may help decide what to review with a qualified
     professional."

4. No diagnostic overreach:
   - Do not say the person has a disease or will develop one.
   - Do not convert common polymorphisms, GWAS signals, or PGx associations into
     deterministic clinical findings.
   - Keep "what this could point to" separate from "what is proven".

5. Condense:
   - Prefer fewer, richer sections over many short repetitive warnings.
   - Do not restate the same limitation more than once unless it is directly
     relevant to a specific axis.
   - The final report should read like a bioinformatician's synthesis with a
     clear evidence appendix, not like a raw gene list.
"""


def schema_v2() -> dict:
    confidence_enum = ["High", "Moderate", "Low", "Conflicting"]
    axis_strength = ["Strong", "Moderate", "Weak", "Conflicting"]
    return {
        "type": "object",
        "additionalProperties": False,
        "required": [
            "metadata",
            "biological_verdict",
            "executive_summary",
            "primary_biological_axes",
            "secondary_or_background_signals",
            "sensitive_or_conflicting_findings",
            "repeated_rsid_interpretation",
            "top_findings_for_human_review",
            "what_this_does_not_mean",
            "family_friendly_summary",
            "professional_technical_summary",
            "limitations",
            "next_review_steps",
            "readiness",
            "final_report_text",
        ],
        "properties": {
            "metadata": {
                "type": "object",
                "additionalProperties": False,
                "required": [
                    "language_mode",
                    "audience_mode",
                    "report_depth",
                    "variant_count_observed",
                    "unique_rsid_count",
                    "unique_gene_count",
                    "confidence_distribution",
                ],
                "properties": {
                    "language_mode": {"type": "string", "enum": ["en", "es", "both"]},
                    "audience_mode": {
                        "type": "string",
                        "enum": ["technical", "health_professional", "family", "all"],
                    },
                    "report_depth": {"type": "string", "enum": ["short", "standard", "detailed"]},
                    "variant_count_observed": {"type": "integer"},
                    "unique_rsid_count": {"type": "integer"},
                    "unique_gene_count": {"type": "integer"},
                    "confidence_distribution": {
                        "type": "object",
                        "additionalProperties": False,
                        "required": confidence_enum,
                        "properties": {label: {"type": "integer"} for label in confidence_enum},
                    },
                },
            },
            "biological_verdict": {
                "type": "object",
                "additionalProperties": False,
                "required": [
                    "main_verdict",
                    "most_relevant_systems",
                    "systems_to_contextualize",
                    "why_this_is_useful",
                    "not_a_diagnosis_statement",
                ],
                "properties": {
                    "main_verdict": {"type": "string"},
                    "most_relevant_systems": {"type": "array", "items": {"type": "string"}},
                    "systems_to_contextualize": {"type": "array", "items": {"type": "string"}},
                    "why_this_is_useful": {"type": "string"},
                    "not_a_diagnosis_statement": {"type": "string"},
                },
            },
            "executive_summary": {
                "type": "object",
                "additionalProperties": False,
                "required": ["summary", "key_takeaway", "most_important_axes", "main_caution"],
                "properties": {
                    "summary": {"type": "string"},
                    "key_takeaway": {"type": "string"},
                    "most_important_axes": {"type": "array", "items": {"type": "string"}},
                    "main_caution": {"type": "string"},
                },
            },
            "primary_biological_axes": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "required": [
                        "axis_name",
                        "axis_strength",
                        "interpretive_weight",
                        "what_may_be_modulated",
                        "plain_language_explanation",
                        "technical_explanation",
                        "supporting_genes",
                        "supporting_rsids",
                        "support_type",
                        "independence_note",
                        "strongest_supporting_findings",
                        "weaker_supporting_findings",
                        "what_not_to_conclude",
                        "review_value",
                        "contextual_review_guidance",
                        "cautions",
                    ],
                    "properties": {
                        "axis_name": {"type": "string"},
                        "axis_strength": {"type": "string", "enum": axis_strength},
                        "interpretive_weight": {
                            "type": "string",
                            "enum": ["primary", "secondary", "background"],
                        },
                        "what_may_be_modulated": {"type": "string"},
                        "plain_language_explanation": {"type": "string"},
                        "technical_explanation": {"type": "string"},
                        "supporting_genes": {"type": "array", "items": {"type": "string"}},
                        "supporting_rsids": {"type": "array", "items": {"type": "string"}},
                        "support_type": {
                            "type": "string",
                            "enum": [
                                "multiple_variants_same_gene",
                                "multiple_genes_same_pathway",
                                "functional_variant_with_context",
                                "sensitive_conflicting_finding",
                                "mixed_support",
                            ],
                        },
                        "independence_note": {"type": "string"},
                        "strongest_supporting_findings": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "required": ["gene", "rsID", "reason", "individual_confidence_level"],
                                "properties": {
                                    "gene": {"type": "string"},
                                    "rsID": {"type": "string"},
                                    "reason": {"type": "string"},
                                    "individual_confidence_level": {"type": "string", "enum": confidence_enum},
                                },
                            },
                        },
                        "weaker_supporting_findings": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": False,
                                "required": ["gene", "rsID", "reason", "individual_confidence_level"],
                                "properties": {
                                    "gene": {"type": "string"},
                                    "rsID": {"type": "string"},
                                    "reason": {"type": "string"},
                                    "individual_confidence_level": {"type": "string", "enum": confidence_enum},
                                },
                            },
                        },
                        "what_not_to_conclude": {"type": "string"},
                        "review_value": {"type": "string"},
                        "contextual_review_guidance": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": [
                                "possible_review_domain",
                                "when_it_may_be_relevant",
                                "who_could_review_it",
                                "what_it_should_clarify",
                                "what_not_to_infer",
                            ],
                            "properties": {
                                "possible_review_domain": {"type": "string"},
                                "when_it_may_be_relevant": {"type": "string"},
                                "who_could_review_it": {"type": "string"},
                                "what_it_should_clarify": {"type": "string"},
                                "what_not_to_infer": {"type": "string"},
                            },
                        },
                        "cautions": {"type": "string"},
                    },
                },
            },
            "secondary_or_background_signals": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "required": ["system_or_category", "summary", "reason_it_is_secondary", "supporting_genes", "supporting_rsids"],
                    "properties": {
                        "system_or_category": {"type": "string"},
                        "summary": {"type": "string"},
                        "reason_it_is_secondary": {"type": "string"},
                        "supporting_genes": {"type": "array", "items": {"type": "string"}},
                        "supporting_rsids": {"type": "array", "items": {"type": "string"}},
                    },
                },
            },
            "sensitive_or_conflicting_findings": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "required": ["gene", "rsID", "why_it_matters", "why_it_is_sensitive_or_conflicting", "how_to_communicate_it", "recommended_review"],
                    "properties": {
                        "gene": {"type": "string"},
                        "rsID": {"type": "string"},
                        "why_it_matters": {"type": "string"},
                        "why_it_is_sensitive_or_conflicting": {"type": "string"},
                        "how_to_communicate_it": {"type": "string"},
                        "recommended_review": {"type": "string"},
                    },
                },
            },
            "repeated_rsid_interpretation": {
                "type": "object",
                "additionalProperties": False,
                "required": ["summary", "repeated_rsids", "how_to_count_them", "why_this_matters"],
                "properties": {
                    "summary": {"type": "string"},
                    "repeated_rsids": {"type": "array", "items": {"type": "string"}},
                    "how_to_count_them": {"type": "string"},
                    "why_this_matters": {"type": "string"},
                },
            },
            "top_findings_for_human_review": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "required": ["gene", "rsID", "review_priority", "reason_for_review", "what_review_should_clarify"],
                    "properties": {
                        "gene": {"type": "string"},
                        "rsID": {"type": "string"},
                        "review_priority": {"type": "string", "enum": ["high", "medium", "low"]},
                        "reason_for_review": {"type": "string"},
                        "what_review_should_clarify": {"type": "string"},
                    },
                },
            },
            "what_this_does_not_mean": {
                "type": "object",
                "additionalProperties": False,
                "required": ["summary", "not_diagnostic", "not_predictive_by_itself", "not_treatment_guidance"],
                "properties": {
                    "summary": {"type": "string"},
                    "not_diagnostic": {"type": "boolean"},
                    "not_predictive_by_itself": {"type": "boolean"},
                    "not_treatment_guidance": {"type": "boolean"},
                },
            },
            "family_friendly_summary": {
                "type": "object",
                "additionalProperties": False,
                "required": [
                    "short_summary",
                    "main_systems_to_pay_attention_to",
                    "simple_explanation",
                    "what_should_not_be_overinterpreted",
                    "what_may_be_worth_discussing_with_a_professional",
                ],
                "properties": {
                    "short_summary": {"type": "string"},
                    "main_systems_to_pay_attention_to": {"type": "array", "items": {"type": "string"}},
                    "simple_explanation": {"type": "string"},
                    "what_should_not_be_overinterpreted": {"type": "string"},
                    "what_may_be_worth_discussing_with_a_professional": {"type": "string"},
                },
            },
            "professional_technical_summary": {
                "type": "object",
                "additionalProperties": False,
                "required": [
                    "methodological_summary",
                    "main_biological_interpretation",
                    "data_quality_comment",
                    "variant_prioritization_comment",
                    "interpretation_constraints",
                ],
                "properties": {
                    "methodological_summary": {"type": "string"},
                    "main_biological_interpretation": {"type": "string"},
                    "data_quality_comment": {"type": "string"},
                    "variant_prioritization_comment": {"type": "string"},
                    "interpretation_constraints": {"type": "string"},
                },
            },
            "limitations": {"type": "array", "items": {"type": "string"}},
            "next_review_steps": {"type": "array", "items": {"type": "string"}},
            "readiness": {
                "type": "object",
                "additionalProperties": False,
                "required": ["ready_for_informational_report", "ready_for_clinical_use", "overall_readiness", "explanation"],
                "properties": {
                    "ready_for_informational_report": {"type": "boolean"},
                    "ready_for_clinical_use": {"type": "boolean"},
                    "overall_readiness": {
                        "type": "string",
                        "enum": ["ready_for_informational_use", "ready_with_minor_review", "needs_technical_review"],
                    },
                    "explanation": {"type": "string"},
                },
            },
            "final_report_text": {
                "type": "object",
                "additionalProperties": False,
                "required": ["title", "report"],
                "properties": {
                    "title": {"type": "string"},
                    "report": {"type": "string"},
                },
            },
        },
    }


def call_openai(prompt: str, payload: dict, model: str, timeout: int, user_instruction: str | None = None) -> dict:
    api_key = os.environ.get("HEAL_OPENAI_API_KEY") or os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("HEAL_OPENAI_API_KEY or OPENAI_API_KEY is required.")
    body = {
        "model": model,
        "input": [
            {"role": "system", "content": prompt},
            {
                "role": "user",
                "content": (
                    (user_instruction or "Generate the final HEAL biological synthesis report from this payload. ")
                    + "Return JSON matching the schema only.\n\n"
                    + json.dumps(payload, ensure_ascii=False, indent=2)
                ),
            },
        ],
        "text": {
            "format": {
                "type": "json_schema",
                "name": "heal_global_interpretation_v2",
                "strict": True,
                "schema": schema_v2(),
            }
        },
    }
    request = urllib.request.Request(
        OPENAI_RESPONSES_URL,
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            parsed = json.loads(response.read().decode("utf-8", errors="replace"))
    except urllib.error.HTTPError as error:
        detail = error.read().decode("utf-8", errors="replace")[:2000]
        raise RuntimeError(f"OpenAI API http_{error.code}: {detail}") from error
    return json.loads(output_text_from_response(parsed))


def translation_prompt(target_language: str) -> str:
    return f"""
You are a strict medical-report translation assistant.

Translate the provided HEAL report JSON into {target_language}.

Critical requirements:
- Preserve the exact JSON structure and all keys required by the schema.
- Preserve the number, order, and grouping of all biological axes.
- Preserve the number, order, and grouping of secondary findings, sensitive findings, limitations, and next steps.
- Do not add, remove, merge, split, reprioritize, or reinterpret axes.
- Do not change axis_strength, interpretive_weight, support_type, confidence labels, booleans, counts, gene symbols, rsIDs, dates, model names, or metadata counts.
- Translate human-readable prose only.
- Keep gene symbols and rsIDs exactly unchanged.
- Keep technical labels such as High, Moderate, Low, Conflicting exactly unchanged.
- Keep the same clinical caution level and the same contextual review guidance strength.
- The translated report must read naturally, but it must remain a translation of the source report, not a new analysis.
"""


def translate_report(report: dict, *, target_language_mode: str, model: str, timeout: int) -> dict:
    payload = {
        "target_language_mode": target_language_mode,
        "source_report": report,
    }
    translated = call_openai(
        translation_prompt("English" if target_language_mode == "en" else "Spanish"),
        payload,
        model,
        timeout,
        user_instruction="Translate this HEAL report JSON while preserving its exact structure. ",
    )
    translated.setdefault("metadata", {})
    translated["metadata"]["language_mode"] = target_language_mode
    translated["metadata"]["generated_at"] = report.get("metadata", {}).get("generated_at", utc_now())
    translated["metadata"]["model"] = report.get("metadata", {}).get("model", model)
    translated["metadata"]["translation_model"] = model
    translated["metadata"]["translation_source_language_mode"] = report.get("metadata", {}).get("language_mode", "es")
    return translated


def add_heading(document, text: str, level: int = 1):
    document.add_heading(str(text or "").strip(), level=level)


def add_para(document, text: str):
    clean = normalize_text(text)
    if not clean:
        return
    for paragraph in clean.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            document.add_paragraph(paragraph)


def normalize_text(text: str) -> str:
    clean = str(text or "").strip()
    clean = clean.replace("\\n", "\n")
    clean = clean.replace("asiociativos", "asociativos")
    clean = clean.replace("asociativa/os", "asociativos")
    return clean


def labels_for(language_mode: str) -> dict[str, str]:
    if language_mode == "en":
        return {
            "observed_variants": "Observed variants",
            "unique_genes": "Unique genes",
            "unique_rsids": "Unique rsIDs",
            "overview": "1. Case overview",
            "key_takeaway": "Key takeaway: ",
            "axes": "2. Main biological axes",
            "review_branch": "Review implication: ",
            "contextual_guidance": "Contextual review guidance",
            "domain": "Possible domain: ",
            "when": "When it may be relevant: ",
            "who": "Who could review it: ",
            "clarify": "What it should clarify: ",
            "not_infer": "What not to infer: ",
            "not_conclude": "Do not conclude: ",
            "technical_support": "Technical support: ",
            "supporting_genes": "Genes supporting this axis: ",
            "relevant_rsids": "Relevant rsIDs: ",
            "other": "3. Secondary findings, limitations, and review steps",
            "sensitive": "Sensitive or conflicting findings",
            "secondary": "Secondary or background signals",
            "family": "Family-facing summary",
            "technical": "Technical/professional summary",
            "limitations": "Limitations",
            "next_steps": "Next review steps",
            "readiness": "Readiness",
        }
    return {
        "observed_variants": "Variantes observadas",
        "unique_genes": "Genes unicos",
        "unique_rsids": "rsIDs unicos",
        "overview": "1. Panorama general del caso",
        "key_takeaway": "Idea central: ",
        "axes": "2. Ejes biologicos principales",
        "review_branch": "Ramificacion de revision: ",
        "contextual_guidance": "Guia de revision contextual",
        "domain": "Dominio posible: ",
        "when": "Cuando podria ser relevante: ",
        "who": "Quien podria revisarlo: ",
        "clarify": "Que deberia aclarar: ",
        "not_infer": "Que no inferir: ",
        "not_conclude": "No concluir: ",
        "technical_support": "Soporte tecnico: ",
        "supporting_genes": "Genes que soportan este eje: ",
        "relevant_rsids": "rsIDs relevantes: ",
        "other": "3. Hallazgos secundarios, limitaciones y pasos de revision",
        "sensitive": "Hallazgos sensibles o conflictivos",
        "secondary": "Senales secundarias o de fondo",
        "family": "Resumen para familia",
        "technical": "Resumen tecnico/profesional",
        "limitations": "Limitaciones",
        "next_steps": "Proximos pasos de revision",
        "readiness": "Readiness",
    }


def render_docx(report: dict, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    add_heading(document, report["final_report_text"]["title"], 0)
    metadata = report["metadata"]
    labels = labels_for(metadata.get("language_mode", "es"))
    add_para(
        document,
        f"{labels['observed_variants']}: {metadata['variant_count_observed']} | "
        f"{labels['unique_genes']}: {metadata['unique_gene_count']} | "
        f"{labels['unique_rsids']}: {metadata['unique_rsid_count']}",
    )

    add_heading(document, labels["overview"], 1)
    add_para(document, labels["key_takeaway"] + report["executive_summary"]["key_takeaway"])
    add_para(document, report["biological_verdict"]["main_verdict"])
    add_para(document, report["biological_verdict"]["why_this_is_useful"])

    add_heading(document, labels["axes"], 1)
    for axis in report["primary_biological_axes"]:
        add_heading(document, f"{axis['axis_name']} ({axis['axis_strength']})", 2)
        add_para(document, axis["what_may_be_modulated"])
        add_para(document, axis["plain_language_explanation"])
        add_para(document, labels["review_branch"] + axis["review_value"])
        guidance = axis["contextual_review_guidance"]
        add_heading(document, labels["contextual_guidance"], 3)
        add_para(document, labels["domain"] + guidance["possible_review_domain"])
        add_para(document, labels["when"] + guidance["when_it_may_be_relevant"])
        add_para(document, labels["who"] + guidance["who_could_review_it"])
        add_para(document, labels["clarify"] + guidance["what_it_should_clarify"])
        add_para(document, labels["not_infer"] + guidance["what_not_to_infer"])
        add_para(document, labels["not_conclude"] + axis["what_not_to_conclude"])
        add_para(document, labels["technical_support"] + axis["independence_note"])
        add_para(document, labels["supporting_genes"] + ", ".join(axis["supporting_genes"]))
        add_para(document, labels["relevant_rsids"] + ", ".join(axis["supporting_rsids"]))

    add_heading(document, labels["other"], 1)
    add_heading(document, labels["sensitive"], 2)
    for item in report["sensitive_or_conflicting_findings"]:
        add_heading(document, f"{item['gene']} {item['rsID']}", 3)
        add_para(document, item["why_it_matters"])
        add_para(document, item["why_it_is_sensitive_or_conflicting"])
        add_para(document, item["recommended_review"])

    add_heading(document, labels["secondary"], 2)
    for item in report["secondary_or_background_signals"]:
        add_heading(document, item["system_or_category"], 3)
        add_para(document, item["summary"])
        add_para(document, item["reason_it_is_secondary"])

    add_heading(document, labels["family"], 2)
    family = report["family_friendly_summary"]
    add_para(document, family["short_summary"])
    add_para(document, family["simple_explanation"])
    add_para(document, family["what_should_not_be_overinterpreted"])
    add_para(document, family["what_may_be_worth_discussing_with_a_professional"])

    add_heading(document, labels["technical"], 2)
    technical = report["professional_technical_summary"]
    for key in [
        "methodological_summary",
        "main_biological_interpretation",
        "data_quality_comment",
        "variant_prioritization_comment",
        "interpretation_constraints",
    ]:
        add_para(document, technical[key])

    add_heading(document, labels["limitations"], 2)
    for item in report["limitations"]:
        document.add_paragraph(item, style="List Bullet")

    add_heading(document, labels["next_steps"], 2)
    for item in report["next_review_steps"]:
        document.add_paragraph(item, style="List Number")

    add_heading(document, labels["readiness"], 2)
    readiness = report["readiness"]
    add_para(document, readiness["overall_readiness"])
    add_para(document, readiness["explanation"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-csv", required=True)
    parser.add_argument("--prompt", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--model", default=os.environ.get("HEAL_LLM2_FULL_MODEL") or "gpt-5.2")
    parser.add_argument("--translation-model", default=os.environ.get("HEAL_LLM2_TRANSLATION_MODEL") or os.environ.get("HEAL_LLM2_FULL_MODEL") or "gpt-5.2")
    parser.add_argument("--language-mode", default="es")
    parser.add_argument("--audience-mode", default="family")
    parser.add_argument("--report-depth", default="detailed")
    parser.add_argument("--timeout-seconds", type=int, default=360)
    args = parser.parse_args()

    config_root_raw = os.environ.get("HEAL_CONFIG_ROOT", "").strip()
    if config_root_raw:
        load_env_file(Path(config_root_raw).expanduser() / "heal-vcf-api.env")
    global_mod = load_global_module()
    rows = read_csv(Path(args.input_csv))
    requested_language_mode = args.language_mode
    base_language_mode = "es" if requested_language_mode in {"en", "both"} else requested_language_mode
    payload = {
        "language_mode": base_language_mode,
        "audience_mode": args.audience_mode,
        "report_depth": args.report_depth,
        "ascii_only": True,
        "project_context": global_mod.project_context(),
        "deterministic_summary": global_mod.build_deterministic_summary(rows),
        "variant_interpretations": [global_mod.compact_variant(row, base_language_mode) for row in rows],
    }

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "llm2_v2_payload.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    prompt = Path(args.prompt).read_text(encoding="utf-8") + report_style_guidance()
    base_report = call_openai(prompt, payload, args.model, args.timeout_seconds)
    base_report["metadata"]["generated_at"] = utc_now()
    base_report["metadata"]["model"] = args.model

    outputs = {}
    base_json_path = output_dir / "heal-global-interpretation-v2-es-source.json"
    base_docx_path = output_dir / "heal-final-report-v2-es-source.docx"
    base_json_path.write_text(json.dumps(base_report, ensure_ascii=False, indent=2), encoding="utf-8")
    render_docx(base_report, base_docx_path)
    outputs["es_source"] = {"json": str(base_json_path), "docx": str(base_docx_path)}

    if requested_language_mode == "en":
        report = translate_report(
            base_report,
            target_language_mode="en",
            model=args.translation_model,
            timeout=args.timeout_seconds,
        )
        json_path = output_dir / "heal-global-interpretation-v2-en.json"
        docx_path = output_dir / "heal-final-report-v2-en.docx"
    else:
        report = base_report
        json_path = output_dir / "heal-global-interpretation-v2.json"
        docx_path = output_dir / "heal-final-report-v2.docx"

    json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    render_docx(report, docx_path)
    outputs["requested"] = {"json": str(json_path), "docx": str(docx_path)}

    if requested_language_mode == "both":
        english_report = translate_report(
            base_report,
            target_language_mode="en",
            model=args.translation_model,
            timeout=args.timeout_seconds,
        )
        english_json_path = output_dir / "heal-global-interpretation-v2-en.json"
        english_docx_path = output_dir / "heal-final-report-v2-en.docx"
        english_json_path.write_text(json.dumps(english_report, ensure_ascii=False, indent=2), encoding="utf-8")
        render_docx(english_report, english_docx_path)
        outputs["en"] = {"json": str(english_json_path), "docx": str(english_docx_path)}

    print(json.dumps({"outputs": outputs, "model": args.model, "translation_model": args.translation_model}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
